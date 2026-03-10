#!/usr/bin/env python3
"""
MD to DOCX Converter (v2 - Modern Design)
범용 마크다운 → DOCX 변환 스크립트

사용법:
    python convert.py file.md                    # 단일 파일
    python convert.py folder/                    # 폴더 내 모든 .md
    python convert.py folder/ --pattern "1차_*"  # 패턴 지정
    python convert.py file.md --output-dir out/  # 출력 폴더 지정
    python convert.py file.md --pdf              # PDF도 함께 생성
"""

import argparse
import re
import os
import sys
import subprocess
from pathlib import Path
from typing import List, Tuple, Optional
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml


# ──────────────────────────────────────────────
# Color System (Indigo-based, refined)
# ──────────────────────────────────────────────
PRIMARY = RGBColor(79, 70, 229)       # Indigo 600
PRIMARY_LIGHT = RGBColor(99, 102, 241)  # Indigo 500
TEXT_COLOR = RGBColor(30, 41, 59)      # Slate 800
TEXT_SECONDARY = RGBColor(100, 116, 139)  # Slate 500
BORDER_COLOR = 'E2E8F0'               # Slate 200 (hex string)
BG_ALT = 'F8FAFC'                     # Slate 50 (hex string)
TABLE_HEADER_BG = '4F46E5'            # Indigo 600 (hex string)
QUOTE_BG = 'EEF2FF'                   # Indigo 50 (hex string)
QUOTE_BAR = '4F46E5'                  # Indigo 600 (hex string)
CODE_BG = 'F1F5F9'                    # Slate 100 (hex string)
WHITE = RGBColor(255, 255, 255)

# Font settings
FONT_FAMILY = 'Pretendard'
FONT_FAMILY_CODE = 'Monaco'


# ──────────────────────────────────────────────
# Helper Functions
# ──────────────────────────────────────────────

def set_korean_font(run, font_name: str = FONT_FAMILY):
    """한글 폰트 설정 (동아시아 폰트 지정)"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_shading(cell, color_hex: str):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """셀 테두리 설정. 각 인자: (size, color, style) 또는 None (테두리 없음)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')

    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{edge}')
        if val is None:
            border.set(qn('w:val'), 'nil')
        else:
            sz, color, style = val
            border.set(qn('w:val'), style)
            border.set(qn('w:sz'), str(sz))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
        tcBorders.append(border)

    # Remove existing tcBorders if any
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    """셀 내부 마진 설정 (단위: twips)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)

    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def add_paragraph_border(paragraph, side='bottom', size='6', color='E2E8F0', style='single'):
    """문단에 테두리 추가"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), style)
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)
    pBdr.append(border)
    pPr.append(pBdr)


def add_page_number(paragraph):
    """현재 페이지 번호 필드 삽입"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)


# ──────────────────────────────────────────────
# Document Setup
# ──────────────────────────────────────────────

def create_styled_document() -> Document:
    """스타일이 적용된 문서 생성"""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = FONT_FAMILY
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT_COLOR
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = FONT_FAMILY
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = TEXT_COLOR
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.2
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = FONT_FAMILY
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = PRIMARY
    h2.paragraph_format.space_before = Pt(28)
    h2.paragraph_format.space_after = Pt(10)
    h2.paragraph_format.line_spacing = 1.3
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = FONT_FAMILY
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = TEXT_COLOR
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.3
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = FONT_FAMILY
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.color.rgb = TEXT_SECONDARY
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(4)
    h4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_FAMILY)

    return doc


def setup_header_footer(doc: Document, branding: str = "Company", footer_text: str = "example.com"):
    """실제 Word 헤더/푸터 설정"""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(branding)
        set_korean_font(run, FONT_FAMILY)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(203, 213, 225)  # Slate 300

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False

        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number
        add_page_number(fp)

        run_sep = fp.add_run('  |  ')
        set_korean_font(run_sep, FONT_FAMILY)
        run_sep.font.size = Pt(8)
        run_sep.font.color.rgb = RGBColor(203, 213, 225)

        run_ft = fp.add_run(footer_text)
        set_korean_font(run_ft, FONT_FAMILY)
        run_ft.font.size = Pt(8)
        run_ft.font.color.rgb = RGBColor(203, 213, 225)


# ──────────────────────────────────────────────
# Title Section
# ──────────────────────────────────────────────

def add_title_section(doc: Document, title: str, branding: str = "Your Organization"):
    """모던 타이틀 섹션"""
    # Branding text (small)
    brand_p = doc.add_paragraph()
    brand_p.paragraph_format.space_before = Pt(20)
    brand_p.paragraph_format.space_after = Pt(4)
    run = brand_p.add_run(branding)
    set_korean_font(run, FONT_FAMILY)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    # Letter spacing
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), '60')
    rPr.append(spacing)

    # Main title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run(title)
    set_korean_font(run, FONT_FAMILY)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = TEXT_COLOR

    # Indigo separator line
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(16)
    add_paragraph_border(sep, 'bottom', '18', '4F46E5', 'single')


# ──────────────────────────────────────────────
# Content Elements
# ──────────────────────────────────────────────

def add_blockquote(doc: Document, lines: List[str]):
    """인용 블록 - 좌측 Indigo 바 + 연한 배경"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)
    set_cell_shading(cell, QUOTE_BG)
    set_cell_margins(cell, top=60, bottom=60, left=160, right=120)

    # Left accent bar + other borders
    set_cell_borders(
        cell,
        left=(18, QUOTE_BAR, 'single'),
        top=(2, BORDER_COLOR, 'single'),
        bottom=(2, BORDER_COLOR, 'single'),
        right=(2, BORDER_COLOR, 'single'),
    )

    # First line uses existing paragraph
    first = True
    for line_text in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()

        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.4

        process_inline_formatting(p, line_text, font_size=Pt(10))

    # Spacing after
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, code: str, language: str = ''):
    """코드 블록 - 둥근 느낌의 회색 배경"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=80, bottom=80, left=160, right=120)
    set_cell_borders(
        cell,
        top=(4, BORDER_COLOR, 'single'),
        bottom=(4, BORDER_COLOR, 'single'),
        left=(4, BORDER_COLOR, 'single'),
        right=(4, BORDER_COLOR, 'single'),
    )

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.3

    run = p.add_run(code.strip())
    run.font.name = FONT_FAMILY_CODE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), FONT_FAMILY_CODE)
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_COLOR

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def add_styled_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """모던 테이블 - 수직선 없음, 수평선만, 줄무늬 행"""
    if not headers or not rows:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Remove default table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)

    # Table cell margins (global)
    tblCellMar = OxmlElement('w:tblCellMar')
    for edge, val in [('top', '60'), ('bottom', '60'), ('start', '100'), ('end', '100')]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), val)
        el.set(qn('w:type'), 'dxa')
        tblCellMar.append(el)
    existing_mar = tblPr.find(qn('w:tblCellMar'))
    if existing_mar is not None:
        tblPr.remove(existing_mar)
    tblPr.append(tblCellMar)

    # ── Header row ──
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, TABLE_HEADER_BG)

        # Header borders: bottom thick indigo line, no vertical
        set_cell_borders(
            cell,
            top=None,
            bottom=(8, '4F46E5', 'single'),
            left=None,
            right=None,
        )

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(header_text.strip())
        set_korean_font(run, FONT_FAMILY)
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = WHITE

    # ── Data rows ──
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        is_last = (row_idx == len(rows) - 1)

        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
            cell = row.cells[col_idx]

            # Zebra striping
            if row_idx % 2 == 1:
                set_cell_shading(cell, BG_ALT)

            # Horizontal borders only
            bottom_border = (4, BORDER_COLOR, 'single') if not is_last else (6, BORDER_COLOR, 'single')
            set_cell_borders(
                cell,
                top=None,
                bottom=bottom_border,
                left=None,
                right=None,
            )

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            _render_cell_text(p, cell_text)

    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def _render_cell_text(paragraph, cell_text: str):
    """테이블 셀 텍스트 렌더링 (볼드, 링크, 코드 등)"""
    # Link handling
    link_match = re.search(r'\[([^\]]+)\]\(([^)]+)\)', cell_text)
    if link_match:
        before = cell_text[:link_match.start()]
        link_text = link_match.group(1)
        after = cell_text[link_match.end():]

        if before:
            run = paragraph.add_run(before)
            set_korean_font(run, FONT_FAMILY)
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT_COLOR

        run = paragraph.add_run(link_text)
        set_korean_font(run, FONT_FAMILY)
        run.font.size = Pt(9.5)
        run.font.color.rgb = PRIMARY
        run.font.underline = True

        if after:
            run = paragraph.add_run(after)
            set_korean_font(run, FONT_FAMILY)
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT_COLOR
        return

    # Inline formatting (bold, code, etc.)
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    parts = re.split(pattern, cell_text)

    for i, part in enumerate(parts):
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_korean_font(run, FONT_FAMILY)
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = TEXT_COLOR
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_FAMILY_CODE
            run.font.size = Pt(8.5)
            run.font.color.rgb = PRIMARY
        else:
            run = paragraph.add_run(part)
            set_korean_font(run, FONT_FAMILY)
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT_COLOR


# ──────────────────────────────────────────────
# Inline Formatting
# ──────────────────────────────────────────────

def process_inline_formatting(paragraph, text: str, font_size=Pt(10.5)):
    """인라인 서식 처리 (굵은 글씨, 코드, 링크)"""
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_korean_font(run, FONT_FAMILY)
            run.font.size = font_size
            run.font.bold = True
            run.font.color.rgb = TEXT_COLOR
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_FAMILY_CODE
            run.font.size = Pt(9)
            run.font.color.rgb = PRIMARY
        elif part.startswith('['):
            match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if match:
                run = paragraph.add_run(match.group(1))
                set_korean_font(run, FONT_FAMILY)
                run.font.size = font_size
                run.font.color.rgb = PRIMARY
                run.font.underline = True
        else:
            run = paragraph.add_run(part)
            set_korean_font(run, FONT_FAMILY)
            run.font.size = font_size
            run.font.color.rgb = TEXT_COLOR


# ──────────────────────────────────────────────
# Markdown Parsing
# ──────────────────────────────────────────────

def parse_table(lines: List[str], start_idx: int) -> Tuple[List[str], List[List[str]], int]:
    """테이블 파싱"""
    header_line = lines[start_idx]
    headers = [h.strip() for h in header_line.strip('|').split('|')]

    # Skip separator
    idx = start_idx + 2

    rows = []
    while idx < len(lines) and lines[idx].strip().startswith('|'):
        row_line = lines[idx]
        row_data = [cell.strip() for cell in row_line.strip('|').split('|')]
        rows.append(row_data)
        idx += 1

    return headers, rows, idx


def detect_doc_type(filename: str) -> str:
    """파일명에서 문서 유형 추출"""
    mapping = {
        '제안서': '통합 제안서',
        '일정표': '일정표',
        '견적서': '비용 견적서',
        '비용견적서': '비용 견적서',
        'OT팀빌딩': 'OT & 팀빌딩',
        '설치가이드': '설치 가이드',
        '핵심개념': '핵심 개념 정리',
        '참고링크': '참고 링크 모음',
        '과제안내': '과제 안내',
    }
    for key, val in mapping.items():
        if key in filename:
            return val
    return '문서'


# ──────────────────────────────────────────────
# Main Conversion
# ──────────────────────────────────────────────

def convert_markdown_to_docx(
    md_path: str,
    output_path: str,
    branding: str = "Company",
    footer_text: str = "example.com"
):
    """마크다운을 DOCX로 변환 (v2 - Modern Design)"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = create_styled_document()

    # Header/Footer
    setup_header_footer(doc, branding, footer_text)

    # Extract title from first H1
    title = ''
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title = stripped.lstrip('#').strip()
            break

    # Title section
    if title:
        add_title_section(doc, title, "Your Organization")

    # Process metadata lines (bold key-value pairs at top)
    i = 0
    in_code_block = False
    code_content = []
    code_language = ''
    found_first_heading = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code block ──
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_content), code_language)
                in_code_block = False
                code_content = []
                code_language = ''
            else:
                in_code_block = True
                code_language = stripped[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # ── Blank line ──
        if not stripped:
            i += 1
            continue

        # ── Horizontal rule ──
        if stripped in ('---', '***', '___'):
            if found_first_heading:
                hr = doc.add_paragraph()
                hr.paragraph_format.space_before = Pt(6)
                hr.paragraph_format.space_after = Pt(6)
                add_paragraph_border(hr, 'bottom', '4', BORDER_COLOR, 'single')
            i += 1
            continue

        # ── Headings ──
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()

            if level == 1:
                if not found_first_heading:
                    found_first_heading = True
                    # Skip H1 since we already rendered it in title section
                    i += 1
                    continue
                doc.add_paragraph(text, style='Heading 1')
            elif level == 2:
                found_first_heading = True
                doc.add_paragraph(text, style='Heading 2')
            elif level == 3:
                doc.add_paragraph(text, style='Heading 3')
            elif level == 4:
                doc.add_paragraph(text, style='Heading 4')
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_korean_font(run, FONT_FAMILY)
                run.font.bold = True
            i += 1
            continue

        # ── Blockquote (multiline support) ──
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                q_text = lines[i].strip()[1:].strip()
                if q_text:
                    quote_lines.append(q_text)
                i += 1
            if quote_lines:
                add_blockquote(doc, quote_lines)
            continue

        # ── Table ──
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, next_idx = parse_table(lines, i)
            add_styled_table(doc, headers, rows)
            i = next_idx
            continue

        # ── Checkbox list ──
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            checked = stripped.startswith('- [x]')
            text = stripped[6:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            checkbox = '  ' if checked else '  '
            run = p.add_run(checkbox)
            set_korean_font(run, FONT_FAMILY)
            run.font.size = Pt(10.5)
            process_inline_formatting(p, text)
            i += 1
            continue

        # ── Unordered list ──
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            process_inline_formatting(p, stripped[2:])
            i += 1
            continue

        # ── Ordered list ──
        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            text = re.sub(r'^\d+\.\s', '', stripped)
            process_inline_formatting(p, text)
            i += 1
            continue

        # ── Metadata lines (bold key: value at start of doc) ──
        if not found_first_heading and re.match(r'^\*\*[^*]+\*\*', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            process_inline_formatting(p, stripped, font_size=Pt(10))
            i += 1
            continue

        # ── Italic line (footer-like) ──
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            text = stripped.strip('*').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            set_korean_font(run, FONT_FAMILY)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = TEXT_SECONDARY
            i += 1
            continue

        # ── Normal paragraph ──
        p = doc.add_paragraph()
        process_inline_formatting(p, stripped)
        i += 1

    # Save
    doc.save(output_path)
    return output_path


# ──────────────────────────────────────────────
# PDF Conversion
# ──────────────────────────────────────────────

def convert_to_pdf(docx_path: str, output_dir: str = None) -> Optional[str]:
    """DOCX를 PDF로 변환 (LibreOffice headless)"""
    soffice = '/opt/homebrew/bin/soffice'
    if not os.path.exists(soffice):
        # Fallback
        soffice = 'soffice'

    if output_dir is None:
        output_dir = os.path.dirname(docx_path) or '.'

    try:
        result = subprocess.run(
            [soffice, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            pdf_name = Path(docx_path).stem + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_name)
            if os.path.exists(pdf_path):
                return pdf_path
        else:
            print(f"  PDF 변환 오류: {result.stderr.strip()}")
    except FileNotFoundError:
        print("  LibreOffice(soffice) 미설치 - PDF 변환 불가")
    except subprocess.TimeoutExpired:
        print("  PDF 변환 타임아웃 (60s)")

    return None


# ──────────────────────────────────────────────
# File Discovery
# ──────────────────────────────────────────────

def find_md_files(path: str, pattern: Optional[str] = None) -> List[Path]:
    """마크다운 파일 찾기"""
    target = Path(path)

    if target.is_file():
        if target.suffix.lower() == '.md':
            return [target]
        return []

    if target.is_dir():
        if pattern:
            if not pattern.endswith('.md'):
                pattern = f"{pattern}.md" if not pattern.endswith('*') else f"{pattern}.md"
            files = sorted(target.glob(pattern))
            return [f for f in files if f.suffix.lower() == '.md']
        return sorted(target.glob('*.md'))

    return []


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='마크다운 → DOCX 변환 (범용 버전)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
예시:
  %(prog)s file.md                      # 단일 파일 변환
  %(prog)s folder/                      # 폴더 내 모든 .md
  %(prog)s folder/ --pattern "1차_*"    # 패턴 매칭
  %(prog)s file.md --output-dir out/    # 출력 폴더
  %(prog)s file.md --pdf                # PDF도 생성
        """
    )

    parser.add_argument('path', help='마크다운 파일 또는 폴더 경로')
    parser.add_argument('--pattern', '-p', help='파일 패턴 (예: "1차_*")')
    parser.add_argument('--output-dir', '-o', help='출력 폴더')
    parser.add_argument('--branding', '-b', default='Company', help='브랜딩 텍스트')
    parser.add_argument('--footer', '-f', default='example.com', help='푸터 텍스트')
    parser.add_argument('--pdf', action='store_true', help='PDF도 함께 생성')

    args = parser.parse_args()

    md_files = find_md_files(args.path, args.pattern)
    if not md_files:
        print("변환할 마크다운 파일이 없습니다.")
        sys.exit(1)

    print(f"\n=== MD -> DOCX Converter v2 ({args.branding}) ===\n")
    print(f"변환 대상: {len(md_files)}개 파일")
    if args.pdf:
        print("PDF 변환: 활성화\n")
    else:
        print()

    output_dir = Path(args.output_dir) if args.output_dir else None
    if output_dir:
        output_dir.mkdir(parents=True, exist_ok=True)

    converted_docx = []
    converted_pdf = []

    for md_file in md_files:
        if output_dir:
            output_file = output_dir / md_file.with_suffix('.docx').name
        else:
            output_file = md_file.with_suffix('.docx')

        print(f"  DOCX: {md_file.name}")
        try:
            result = convert_markdown_to_docx(
                str(md_file),
                str(output_file),
                branding=args.branding,
                footer_text=args.footer
            )
            converted_docx.append(result)
            print(f"    -> {output_file}")

            if args.pdf:
                pdf_dir = str(output_dir) if output_dir else str(md_file.parent)
                pdf_result = convert_to_pdf(str(output_file), pdf_dir)
                if pdf_result:
                    converted_pdf.append(pdf_result)
                    print(f"    -> {pdf_result}")
        except Exception as e:
            print(f"    ERROR: {e}")
            import traceback
            traceback.print_exc()

    print(f"\n=== 완료: DOCX {len(converted_docx)}/{len(md_files)}", end='')
    if args.pdf:
        print(f", PDF {len(converted_pdf)}/{len(md_files)}", end='')
    print(" ===\n")


if __name__ == '__main__':
    main()
